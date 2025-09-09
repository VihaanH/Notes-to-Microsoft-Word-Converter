from pathlib import Path
from docx import Document
from pdf2image import convert_from_path
from PIL import Image, ImageOps
import numpy as np, cv2, torch, io

from transformers import TrOCRProcessor, VisionEncoderDecoderModel

# CONFIG
USE_EASYOCR = False                          # flip to True to compare EasyOCR
SAVE_DEBUG_LINES = True                      # saves crops to ./debug_lines/
OCR_MODEL_NAME = "microsoft/trocr-large-handwritten"  # or 'trocr-base-handwritten'
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

# OCR engines
if not USE_EASYOCR:
    processor = TrOCRProcessor.from_pretrained(OCR_MODEL_NAME)
    model = VisionEncoderDecoderModel.from_pretrained(OCR_MODEL_NAME).to(DEVICE)
else:
    import easyocr, warnings
    warnings.filterwarnings("ignore", category=UserWarning)
    reader = easyocr.Reader(['en'], gpu=DEVICE == "cuda")

@torch.inference_mode()
def ocr_line(img: Image.Image) -> str:
    """OCR a single cropped line image -> text string."""
    if USE_EASYOCR:
        result = reader.readtext(np.array(img), detail=0, paragraph=True)
        return " ".join(result)
    # ---- TrOCR path ----
    if img.mode != "RGB":
        img = img.convert("RGB")
    if img.height < 48:                      # upscale very thin lines
        s = 48 / img.height
        img = img.resize((int(img.width*s), 48), Image.BICUBIC)

    pix = processor(images=img, return_tensors="pt").pixel_values.to(DEVICE)
    ids  = model.generate(pix)
    return processor.batch_decode(ids, skip_special_tokens=True)[0]

# ────────────────────────────────────────────────────────────────────────
# Pre-processing helpers
# ────────────────────────────────────────────────────────────────────────
def preprocess_for_segmentation(page: Image.Image) -> Image.Image:
    """Invert & boost contrast so handwriting is dark on light."""
    gray = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2GRAY)
    inv  = cv2.bitwise_not(gray)                         # white-on-black → black-on-white
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
    boosted = clahe.apply(inv)
    _, bw = cv2.threshold(boosted, 0, 255,
                          cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return Image.fromarray(bw)

def slice_into_lines(raw_page: Image.Image) -> list[Image.Image]:
    """Detect & crop each text line with generous padding."""
    bw = preprocess_for_segmentation(raw_page)
    bw_np = np.array(bw)

    # connect characters within a line
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT,
                                       (raw_page.width // 30, 4))
    merged = cv2.dilate(bw_np, kernel, iterations=2)

    contours, _ = cv2.findContours(merged, cv2.RETR_EXTERNAL,
                                   cv2.CHAIN_APPROX_SIMPLE)
    pad_y, pad_x = 16, 32
    crops = []
    for cnt in sorted(contours, key=lambda c: cv2.boundingRect(c)[1]):
        x, y, w, h = cv2.boundingRect(cnt)
        if h < 15 or w < 50:
            continue
        crop = raw_page.crop((max(x-pad_x,0), max(y-pad_y,0),
                              min(x+w+pad_x, raw_page.width),
                              min(y+h+pad_y, raw_page.height)))
        crops.append(crop)

    return crops

# Main pipeline
def notes_to_word(src, output="converted_notes.docx"):
    pages = (convert_from_path(src, dpi=300) if src.lower().endswith(".pdf")
             else [Image.open(src)])

    doc = Document(); doc.add_heading("Converted Notes", level=1)

    for i, page in enumerate(pages, 1):
        if len(pages) > 1:
            doc.add_heading(f"Page {i}", level=2)
        for line_img in slice_into_lines(page):
            txt = ocr_line(line_img)
            if txt.strip():
                doc.add_paragraph(txt.strip())

    out = Path(output).resolve(); doc.save(out)
    print("✔ Saved to", out); return out

notes_to_word('/Users/vihaan14/Downloads/Notes.jpg')